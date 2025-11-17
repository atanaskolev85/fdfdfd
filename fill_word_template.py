#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Fill Word template with test data from screenshots
"""
from docx import Document

# Load template
doc = Document("02_785_692_OFFER_APPROVED.docx")

print("Populating Word document with test data...")

# Fill Table 0 (OVERVIEW)
if len(doc.tables) > 0:
    table = doc.tables[0]
    print("Filling OVERVIEW table...")

    # Find and fill each field
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = cells[0].text.strip().lower()

            if 'plant owner' in key:
                cells[1].text = "SE-Plovdiv"
                print("  ✓ Plant owner: SE-Plovdiv")
            elif 'plant code' in key:
                cells[1].text = "BG02"
                print("  ✓ Plant code: BG02")
            elif 'name of the part' in key:
                cells[1].text = "Cover KEO GRECO"
                print("  ✓ Name of the part: Cover KEO GRECO")
            elif 'reference' in key and 'tool' not in key:
                cells[1].text = "GDE51999 GDE51998NVE18753NVE20604"
                print("  ✓ Reference")
            elif 'tool number' in key:
                cells[1].text = "Y1"
                print("  ✓ Tool number: Y1")
            elif 'se inventory' in key:
                cells[1].text = "260000000546"
                print("  ✓ SE inventory number: 260000000546")
            elif 'gotmar inventory' in key:
                cells[1].text = "692"
                print("  ✓ Gotmar inventory: 692")
            elif 'general condition' in key:
                cells[1].text = "Good"
                print("  ✓ General condition: Good")
            elif 'type of service' in key:
                cells[1].text = "Repair"
                print("  ✓ Type of service: Repair")

# Fill Table 1 (DATES)
if len(doc.tables) > 1:
    table = doc.tables[1]
    print("\nFilling DATES table...")

    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = cells[0].text.strip().lower()

            if 'project creation date' in key:
                cells[1].text = "2025-10-22"
                print("  ✓ Project creation date: 2025-10-22")
            elif 'offer creation date' in key:
                cells[1].text = "2025-11-11"
                print("  ✓ Offer creation date: 2025-11-11")
            elif 'approval date' in key:
                cells[1].text = "2025-11-17"
                print("  ✓ Approval date: 2025-11-17")
            elif 'finish of the project (estimated)' in key:
                cells[1].text = "2025-12-29"
                print("  ✓ Finish date (estimated): 2025-12-29")

# Fill Table 2 (TOTAL COST)
if len(doc.tables) > 2:
    table = doc.tables[2]
    print("\nFilling COST table...")

    # Find the "Total cost" row and add value
    for row in table.rows:
        cells = row.cells
        row_text = ' '.join([c.text for c in cells])

        if 'Total cost' in row_text:
            # Put cost in last cell
            if len(cells) > 0:
                # Find first empty cell or last cell
                for i in range(len(cells)-1, -1, -1):
                    if 'Total cost' not in cells[i].text:
                        cells[i].text = "3010"
                        print("  ✓ Total cost: 3010")
                        break

# Save filled document
output_file = "02_785_692_OFFER_APPROVED_FILLED.docx"
doc.save(output_file)

print(f"\n✓ Created filled document: {output_file}")
print("\nNow you can use this file with the converter!")
